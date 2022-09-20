from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
import subprocess
import json
import datetime
import random


def testPrint():
    print("Connected to formatter")

correct_answers = {
    "a1" : "2",
    "a2" : "37",
    "a3" : "2"
}


count =0


answer_lookup = {
    "en": {
        "a1": {
            "1": "Yes",
            "2": "No",
        },
        "a3": {
            "1": "How you see and identify yourself",
            "2": "How others see and identify you",
            "3": "Both", 
            "4": "None of the above"
        },
        "a2": {
            "0": "Anderson",
            "1": "Barret",
            "2": "Buck",
            "3": "Bell",
            "4": "Coats",
            "5": "Cary",
            "6": "Kerry",
            "7": "Duperton",
            "8": "Dipitón", 
            "9": "Dishmey", 
            "10": "Dishmer",
            "11": "Buck",
            "12": "Ejice",
            "13": "Fuchue",
            "14": "Forche",
            "15": "Green", 
            "16": "Hilton", 
            "17": "Henderson",
            "18": "Hopkins",
            "19": "James",
            "20": "Jhonson", 
            "21": "Jones", 
            "22": "King", 
            "23": "Kinxon", 
            "24": "Miller", 
            "25": "Nwes", 
            "26": "Punez", 
            "27": "Paul", 
            "28": "Robinson", 
            "29": "Relmon", 
            "30": "Schod", 
            "31": "Sarry", 
            "32": "Shepherd", 
            "33": "Sapher", 
            "34": "Wingth", 
            "35": "Williams", 
            "36": "Willmore", 
            "37": "None of the above" 
        },
    },
    "es": {
        "a1": {
            "1": "Si",
            "2": "No"
        },
        "a3": {
            "1": "Cómo te ves e identificas",
            "2": "Como otres te ven e identifican",
            "3": "Todo lo anterior",
            "4": "Ninguna de las anteriores"
        },
               
        "a2": {
            "0": "Anderson",
            "1": "Barret",
            "2": "Buck",
            "3": "Bell",
            "4": "Coats",
            "5": "Cary",
            "6": "Kerry",
            "7": "Duperton",
            "8": "Dipitón", 
            "9": "Dishmey", 
            "10": "Dishmer",
            "11": "Buck",
            "12": "Ejice",
            "13": "Fuchue",
            "14": "Forche",
            "15": "Green", 
            "16": "Hilton", 
            "17": "Henderson",
            "18": "Hopkins",
            "19": "James",
            "20": "Jhonson", 
            "21": "Jones", 
            "22": "King", 
            "23": "Kinxon", 
            "24": "Miller", 
            "25": "Nwes", 
            "26": "Punez", 
            "27": "Paul", 
            "28": "Robinson", 
            "29": "Relmon", 
            "30": "Schod", 
            "31": "Sarry", 
            "32": "Shepherd", 
            "33": "Sapher", 
            "34": "Wingth", 
            "35": "Williams", 
            "36": "Willmore", 
            "37": "None of the above" 
        },
    }
}


from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def checkedElement():
    elm = OxmlElement('w:checked')
    elm.set(qn('w:val'),"true")
    return elm

zipcode_data_lookup = {}
country_index_score = {}

# with open('/home/pi/zipcode_data.txt') as json_file:
#     zipcode_data_lookup = json.load(json_file)

with open('country_json_data.json', 'r') as myfile:
    data=myfile.read()

# parse file
country_index_score = json.loads(data)

def score_answers(userInfo, country_score):
    number_correct = 0
    for answer in correct_answers:
        print(answer, userInfo[answer])
        if (userInfo[answer] == correct_answers[answer]):
            print("Correct")
            number_correct +=1

    if country_score <= 23.6:
        number_correct +=1 
    return number_correct


def formatDocument(userInfo):
    print("Starting Custom Print Job")
    print(userInfo)
    doc = Document("/home/pi/CivilResponsesEN.docx")
    lang = userInfo["lang"]
    if userInfo["lang"] == "es":
        doc = Document("/home/pi/CivilResponsesES.docx")


    country_name = country_index_score[userInfo["countryName"]]["country_name"]
    country_score = float(country_index_score[userInfo["countryName"]]["score"])

    print(country_name, country_score)

    styles = doc.styles
    style = styles.add_style('Insertion', WD_STYLE_TYPE.PARAGRAPH)
    style = doc.styles['Insertion']
    font = style.font
    font.name = 'Helvetica'
    font.size = Pt(18)


    num_correct = score_answers(userInfo, country_score)
    qualify_status = "QUALIFY"  if num_correct == 5  else "DISQUALIFY"
    print("QUALIFY : " , qualify_status)
    print("NUM CORRECT ", num_correct)


    print("Creating Document using this info")
    print(userInfo)
    for paragraph in doc.paragraphs:
        # print(paragraph.text)

        if '[DATE]' in paragraph.text:
            paragraph.style = 'Insertion'
            paragraph.text =f"{datetime.datetime.now():%m-%d-%Y}" 

        if '[NAME]' in paragraph.text:
            paragraph.style = 'Insertion'
            paragraph.text = 'Applicant: {}'.format(userInfo["userName"])
            if lang == "es":
                paragraph.text = 'Solicitante: {}'.format(userInfo["userName"])

        if '[QUALIFYSTATUS]' in paragraph.text:
            paragraph.style = 'Insertion'
            paragraph.text = 'Result : {}'.format(qualify_status)
            if lang == "es":
                paragraph.text = "Resultado : No Califica"

        if '[Q1 answer]' in paragraph.text:   
            q1Answer = answer_lookup[lang]["a1"][userInfo["a1"]]

            if lang == "en":
                paragraph.text = "You [{}] know or have heard of the suspects and witnesses.".format(q1Answer)
            else:
                q1Answer = answer_lookup[lang]["a1"][userInfo["a1"]]
                paragraph.text = "Para usted, la historia colonial es [{}] para el presente.".format(q1Answer)

        if '[Q2 answer]' in paragraph.text:
            questionTwoAnswer = answer_lookup[lang]["a2"][userInfo["a2"]]
            paragraph.text = "You are [{}] to a person who is stateless.".format(questionTwoAnswer)
  
            if lang == "es":
                paragraph.text = "Usted es [{}] a una persona apátrida.".format(questionTwoAnswer)

        if '[Q4 answer]' in paragraph.text:
            questionThreeAnswer = answer_lookup[lang]["a3"][userInfo["a3"]]
            print("paragraph text", paragraph.text)
            paragraph.text = "You [{}] that the nation-state is a violent institution.".format(questionThreeAnswer)
            if lang == "es":
                paragraph.text = "Usted está [{}] que el estado-nación es una institución violenta.".format(questionThreeAnswer)

        if '[Q3 answer]' in paragraph.text:
            questionFourAnswer = answer_lookup[lang]["sugarIntake"][userInfo["sugarIntake"]]
            paragraph.text = "Your weekly sugar intake is [{}].  To be an impartial reviewer would not consume any sugar.".format(
                answer_lookup[lang]["sugarIntake"][userInfo["sugarIntake"]])
            if lang == "es":
                paragraph.text = "Su consumo semanal de azúcar es [Q3 answer].".format(questionFourAnswer)

        if '[ANSWER]' in paragraph.text:
            print("Setting Random Value 1")
            paragraph.style = 'Insertion'

            paragraph.text = "Your [{}] nationality ranks {}% in the Quality of Nationality index".format(country_name,country_score)
            if lang == "es":
                paragraph.text = "Su nacionalidad [{}] tiene un índice de {}% en el índice de calidad de la nacionalidad".format(country_name,country_score)


        if '[X out of 5]' in paragraph.text or '[X de 5]' in paragraph.text:
            print("Setting Random Value 1")
            paragraph.style = 'Insertion'

            paragraph.text = "You answered [{} out of 5] questions correctly. To be an impartial reviewer you would have to answer all the questions correctly.".format(num_correct)
            if lang == "es":
                paragraph.text = "Usted obtuvo [{} de 5] correctas. Para ser un evaluador imparcial debe responder correctamente todas las preguntas.".format(num_correct)





    doc_name = '{}.docx'.format(userInfo["userName"])
    doc.save(doc_name)
    print("Formatted And Saved Document with name {}".format(doc_name))
    subprocess.run(["libreoffice", "--headless", "--convert-to",
                   "pdf", "{}.docx".format(userInfo["userName"])])
    subprocess.run(
        ["lp", "-d", "myprinter", "{}.pdf".format(userInfo["userName"])])
    print("Completed Format Of Document")


if __name__ == "__main__":
    try:
        formatDocument({'userName': 'Jsasesdsica', 'userId': 'd63142d7cf6f53a093ebc32ae1448f18', 'a1': '1', 'a2': '2',
                       'a3': '1', 'countryName': "53", 'sugarIntake': '3', 'archivePermission': '1', 'lang': 'es'})
    except Exception as e:
        print(e)
