from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
doc = Document("translations_matt.docx")

testimonies = {
    "Living on an island" : [],
    "What do you know about enslaved Black people in the Dominican Republic?"

}


def extract_english_translations():
    doc = Document("translations_matt.docx")
    print(doc.tables[0].rows[0].cells[0].text)

    for row in doc.tables[0].rows:
        text = row.cells[0].text
        if "Testimony : " in text or "Testimonials :" in text:
            print(  text)


extract_english_translations()