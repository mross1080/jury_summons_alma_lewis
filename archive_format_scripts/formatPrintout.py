from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
import subprocess
import json
import datetime
import random


def testPrint():
    print("Connected to formatter ")

correct_answers = {
    "know_the_witness" : "2",
    "related_to_or_know" : "37",
    "a3" : "2"
}


count =0


answer_lookup = {
    "en": {
        "know_the_witness": {
            "1": "Have",
            "2": "Have not",
        },
        "banner_completion" : ["invader","sovereignty","constitution"],
        "a3": {
            "1": "How you see and identify yourself",
            "2": "How others see and identify you",
            "3": "Both", 
            "4": "None of the above"
        },
        "last_names": {
            "0": "Anderson",
            "1": "Barret",
            "2": "Buck",
            "3": "Bell",
            "4": "Coats",
            "5": "Cary",
            "6": "Kerry",
            "7": "Duperton",
            "8": "Dipitón", 
            "9": "Dishmey", 
            "10": "Dishmer",
            "11": "Buck",
            "12": "Ejice",
            "13": "Fuchue",
            "14": "Forche",
            "15": "Green", 
            "16": "Hilton", 
            "17": "Henderson",
            "18": "Hopkins",
            "19": "James",
            "20": "Jhonson", 
            "21": "Jones", 
            "22": "King", 
            "23": "Kinxon", 
            "24": "Miller", 
            "25": "Nwes", 
            "26": "Punez", 
            "27": "Paul", 
            "28": "Robinson", 
            "29": "Relmon", 
            "30": "Schod", 
            "31": "Sarry", 
            "32": "Shepherd", 
            "33": "Sapher", 
            "34": "Wingth", 
            "35": "Williams", 
            "36": "Willmore", 
            "37": "None of the above" 
        },
    },
    "es": {
        "know_the_witness": {
            "1": "Si",
            "2": "No"
        },
        "a3": {
            "1": "Cómo te ves e identificas",
            "2": "Como otres te ven e identifican",
            "3": "Todo lo anterior",
            "4": "Ninguna de las anteriores"
        },
               
        "last_names": {
            "0": "Anderson",
            "1": "Barret",
            "2": "Buck",
            "3": "Bell",
            "4": "Coats",
            "5": "Cary",
            "6": "Kerry",
            "7": "Duperton",
            "8": "Dipitón", 
            "9": "Dishmey", 
            "10": "Dishmer",
            "11": "Buck",
            "12": "Ejice",
            "13": "Fuchue",
            "14": "Forche",
            "15": "Green", 
            "16": "Hilton", 
            "17": "Henderson",
            "18": "Hopkins",
            "19": "James",
            "20": "Jhonson", 
            "21": "Jones", 
            "22": "King", 
            "23": "Kinxon", 
            "24": "Miller", 
            "25": "Nwes", 
            "26": "Punez", 
            "27": "Paul", 
            "28": "Robinson", 
            "29": "Relmon", 
            "30": "Schod", 
            "31": "Sarry", 
            "32": "Shepherd", 
            "33": "Sapher", 
            "34": "Wingth", 
            "35": "Williams", 
            "36": "Willmore", 
            "37": "None of the above" 
        },
    }
}


from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def checkedElement():
    elm = OxmlElement('w:checked')
    elm.set(qn('w:val'),"true")
    return elm

zipcode_data_lookup = {}
country_index_score = {}

# with open('/home/pi/zipcode_data.txt') as json_file:
#     zipcode_data_lookup = json.load(json_file)

with open('country_json_data.json', 'r') as myfile:
    data=myfile.read()

# parse file
country_index_score = json.loads(data)

def score_answers(userInfo, country_score):
    print("Scoring asnwers ")
    number_correct = 0
    for answer in correct_answers:
        print(answer, userInfo[answer])
        if (userInfo[answer] == correct_answers[answer]):
            print("Correct")
            number_correct +=1

    if country_score <= 23.6:
        number_correct +=1 
    return number_correct


def formatDocument(userInfo):
    print("Starting Custom Print Job")
    print("Incoming data to format print with " , userInfo)
    try: 
        doc = Document("CivilResponseEN.docx")
        lang = userInfo["lang"]
        if userInfo["lang"] == "es":
            doc = Document("CivilResponsesES.docx")

        print("Extrapolating Country name from value") 
        country_name = country_index_score[userInfo["countryName"]]["country_name"]
        print("Country name is ", country_name)
        country_score = float(country_index_score[userInfo["countryName"]]["score"])

        print(country_name, country_score)

        styles = doc.styles
        style = styles.add_style('Insertion', WD_STYLE_TYPE.PARAGRAPH)
        style = doc.styles['Insertion']
        font = style.font
        font.name = 'Helvetica'
        font.size = Pt(18)


        num_correct = score_answers(userInfo, country_score)
        qualify_status = "QUALIFY"  if num_correct == 5  else "DISQUALIFY"
        print("QUALIFY : " , qualify_status)
        print("NUM CORRECT ", num_correct)


        print("Creating Document using this info")
        print(userInfo)
        for paragraph in doc.paragraphs:
            # print(paragraph.text)

            if '[DATE]' in paragraph.text:
                paragraph.style = 'Insertion'
                paragraph.text =f"{datetime.datetime.now():%m-%d-%Y}" 

            if '[NAME]' in paragraph.text:
                paragraph.style = 'Insertion'
                paragraph.text = 'Applicant: {}'.format(userInfo["userName"])
                if lang == "es":
                    paragraph.text = 'Solicitante: {}'.format(userInfo["userName"])

            if '[QUALIFYSTATUS]' in paragraph.text:
                paragraph.style = 'Insertion'
                paragraph.text = 'Result : {}'.format(qualify_status)
                if lang == "es":
                    paragraph.text = "Resultado : No Califica"

            if '[Q1]' in paragraph.text:
                print("Checking banner issues ")
                banner_answers = userInfo["banner_completion"].split(",")
                num_lines_correct = 0
                for index, line in enumerate(banner_answers):
                    print(index, line)
                    print(answer_lookup[lang]["banner_completion"][index])
                    if line.lower() == answer_lookup[lang]["banner_completion"][index]:
                        num_lines_correct+=1
                #q1Answer = answer_lookup[lang]["banner_completion"][userInfo["banner_completion"]]
                # For testing
                q1Answer = num_lines_correct
                if lang == "en":
                    paragraph.text = "You answer [{}] words correctly from the 1965 banner.".format(q1Answer)
                else:
                    q1Answer = answer_lookup[lang]["know_the_witness"][userInfo["know_the_witness"]]
                    paragraph.text = "Usted respondió [{}] palabras correctas en la pancarta del 1965.".format(q1Answer)

                if num_lines_correct == 3:
                    num_correct+=1

            if '[Q2]' in paragraph.text:
                print("Know the witness ", userInfo["know_the_witness"])
                questionTwoAnswer = answer_lookup[lang]["know_the_witness"][userInfo["know_the_witness"]]
                if userInfo["know_the_witness"] == "1":
                    questionTwoAnswer = "do"
                else:
                    questionTwoAnswer = "do not"
                print(userInfo["know_the_witness"], answer_lookup[lang]["know_the_witness"][userInfo["know_the_witness"]])
                paragraph.text = "You [{}] know or have heard of the suspects and witnesses.".format(questionTwoAnswer)
    
                if lang == "es":
                    paragraph.text = "Usted [{}] los testigos o los sospechosos..".format(questionTwoAnswer)
                
                if userInfo["know_the_witness"] == "1":
                    num_correct+=1

            if '[Q3]' in paragraph.text:

                if userInfo["related_to_or_know"] != "":
                    questionThreeAnswer = "are"
                else:
                    questionThreeAnswer = "are not"
                #questionThreeAnswer = answer_lookup[lang]["a3"][userInfo["a3"]]
                #questionThreeAnswer = "are not"
                print("paragraph text", paragraph.text)
                paragraph.text = "You [{}] related or know a descendent of an immigrant from the 1824 migration from the United States to Haiti/Dominican Republic.".format(questionThreeAnswer)
                if lang == "es":
                    paragraph.text = "Usted [{}] ser un familiar o conoce a un descendiente de inmigrante de la migración de 1824 de los Estados Unidos a Haití/República Dominicana.".format(questionThreeAnswer)


            if '[ANSWER]' in paragraph.text:
                paragraph.style = 'Insertion'

                paragraph.text = "[{}] citizenship ranks {}% in the Quality of Nationality index*.".format(country_name,country_score)
                if lang == "es":
                    paragraph.text = "La ciudadanía de [{}] tiene un índice de {}% en el índice de calidad de la nacionalidad*".format(country_name,country_score)


            if '[X out of 4]' in paragraph.text or '[X de 4]' in paragraph.text:
                print("Setting Random Value 1")
                paragraph.style = 'Insertion'

                paragraph.text = "You answered [{} out of 4] questions correctly. To be an impartial reviewer you would have to answer all the questions correctly.".format(num_correct)
                if lang == "es":
                    paragraph.text = "Usted obtuvo [{} de 4] correctas. Para ser un evaluador imparcial debe responder correctamente todas las preguntas.".format(num_correct)





        doc_name = '{}.docx'.format(userInfo["userName"])
        doc.save(doc_name)
        print("Formatted And Saved Document with name {}".format(doc_name))
        subprocess.run(["libreoffice", "--headless", "--convert-to",
                    "pdf", "{}.docx".format(userInfo["userName"])])
        subprocess.run(
            ["lp", "-d", "epson_2023", "{}.pdf".format(userInfo["userName"])])
        print("Completed Format Of Document")
    except Exception as e:
        print("Encountered Error In Printout ", e)


if __name__ == "__main__":
    try:
        # formatDocument({'userName': 'bar', 'userId': '8dfb74f85758a53ae9bb142a8493a341', 'related_to_or_know': '1', 'a3': '1', 'banner_completion': '', 'know_the_witness': '2', 'countryName': '33', 'archivePermission': '1', 'lang': 'en'} )
        formatDocument({'userName': 'aLizania cruz', 'userId': 'b9961849347705fd62b8c00dba1e843c', 'related_to_or_know': '2', 'a3': '1', 'banner_completion': 'Joke,Joke,Joke', 'know_the_witness': '2', 'countryName': '7', 'archivePermission': '1', 'lang': 'en', 'a1': '2'})
    except Exception as e:
        print(e)
